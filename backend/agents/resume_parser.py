import os
import io
import time
from datetime import datetime
from services.s3 import download_file_from_s3
from agents.state import AgentState


def _extract_text_pure_python(file_bytes: bytes, filename: str) -> str:
    """
    Highly optimized 100% Python document parser.
    Uses pdfplumber (best accuracy) -> pypdf -> PyPDF2 for PDFs.
    Uses python-docx for DOCX/DOC files.
    Uses UTF-8 decode for TXT.
    No external API calls or AWS Textract required!
    """
    ext = os.path.splitext(filename)[1].lower() if filename else ".pdf"

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            # Also extract text inside tables (common in resume templates)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in lines:
                            lines.append(cell_text)
            text = "\n".join(lines)
            if text.strip():
                print(f"[ResumeParser] Extracted DOCX via python-docx ({len(text)} chars)")
                return text
        except Exception as e:
            print(f"[ResumeParser] python-docx extraction warning: {e}")

    # PDF Parsing Chain: pdfplumber (highest layout accuracy) -> pypdf -> PyPDF2
    if ext == ".pdf" or file_bytes:
        # Tier 1: pdfplumber (handles multi-column resumes, tables, complex formatting)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                extracted_pages = []
                for page in pdf.pages:
                    txt = page.extract_text(layout=True) or page.extract_text()
                    if txt:
                        extracted_pages.append(txt)
                text = "\n".join(extracted_pages)
                if text.strip():
                    print(f"[ResumeParser] Extracted PDF via pdfplumber ({len(text)} chars)")
                    return text
        except Exception as e:
            print(f"[ResumeParser] pdfplumber fallback warning: {e}")

        # Tier 2: pypdf (fast, robust text stream extraction)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            extracted_pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(extracted_pages)
            if text.strip():
                print(f"[ResumeParser] Extracted PDF via pypdf ({len(text)} chars)")
                return text
        except Exception as e:
            print(f"[ResumeParser] pypdf fallback warning: {e}")

        # Tier 3: PyPDF2 (standard fallback)
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            extracted_pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(extracted_pages)
            if text.strip():
                print(f"[ResumeParser] Extracted PDF via PyPDF2 ({len(text)} chars)")
                return text
        except Exception as e:
            print(f"[ResumeParser] PyPDF2 fallback warning: {e}")

    # If all specialized parsers failed, try raw UTF-8 decode as last resort
    return file_bytes.decode("utf-8", errors="ignore")


def resume_parser_agent(state: AgentState) -> AgentState:
    """
    Agent 1: Downloads resume from S3, extracts text using high-accuracy Python libraries
    (pdfplumber / pypdf / python-docx), then uses Groq LLM to structure into clean JSON.
    Zero external AWS Textract dependency — 100% free-tier compatible.
    """
    from services.llm import llm_call, parse_json_from_llm

    start = time.time()
    s3_key = state.get("s3_key", "")
    filename = s3_key.split("/")[-1] if s3_key else "resume.pdf"
    errors = state.get("errors", [])
    audit_entries = state.get("audit_entries", [])

    if not s3_key:
        errors.append("ResumeParserAgent: s3_key is empty — cannot process resume")
        return {**state, "parsed_resume": {}, "raw_text": "", "errors": errors, "audit_entries": audit_entries}

    try:
        ext = os.path.splitext(filename)[1].lower()

        # Download bytes from S3
        file_bytes = download_file_from_s3(s3_key)
        if not file_bytes:
            raise ValueError(f"S3 file download returned empty bytes for key: {s3_key}")

        # Extract text using pure Python framework pipeline
        raw_text = _extract_text_pure_python(file_bytes, filename)

        if not raw_text.strip():
            raise ValueError("Resume text extraction returned empty text — file may be unreadable or image-only")

        # LLM Structuring
        system_prompt = """You are an expert resume parser. Extract structured information from resume text.
Return ONLY valid JSON with this exact structure:
{
  "skills": ["skill1", "skill2"],
  "experience": [{"title": "Job Title", "company": "Company", "duration": "X years", "description": "responsibilities"}],
  "education": [{"degree": "Degree Name", "institution": "University", "year": "YYYY"}],
  "certifications": ["cert1"],
  "projects": [{"name": "Project", "description": "what it does", "tech_stack": ["tech1"]}],
  "achievements": ["achievement1"]
}
Be thorough. If a field has no data, return an empty array."""

        user_prompt = f"Parse this resume text:\n\n{raw_text[:6000]}"
        response_text, model_used, tokens = llm_call(system_prompt, user_prompt)
        parsed_resume = parse_json_from_llm(response_text)

        duration_ms = int((time.time() - start) * 1000)
        audit_entries.append({
            "agent_name": "ResumeParserAgent",
            "input_summary": f"S3 key: {s3_key}, format: {ext}, text length: {len(raw_text)} chars",
            "output_summary": f"Parsed {len(parsed_resume.get('skills', []))} skills, {len(parsed_resume.get('experience', []))} jobs, {len(parsed_resume.get('projects', []))} projects",
            "reasoning": "Pure Python parser (pdfplumber/pypdf/docx) + Groq LLM structuring",
            "timestamp": datetime.utcnow().isoformat(),
            "duration_ms": duration_ms,
            "llm_model": model_used,
            "tokens_used": tokens,
        })

        return {
            **state,
            "raw_text": raw_text,
            "parsed_resume": parsed_resume,
            "audit_entries": audit_entries,
        }

    except Exception as e:
        err_msg = f"ResumeParserAgent error: {str(e)}"
        errors.append(err_msg)
        print(f"[ResumeParser] FAILED: {err_msg}")
        audit_entries.append({
            "agent_name": "ResumeParserAgent",
            "input_summary": f"S3 key: {s3_key}",
            "output_summary": "FAILED",
            "reasoning": str(e),
            "timestamp": datetime.utcnow().isoformat(),
            "duration_ms": int((time.time() - start) * 1000),
            "llm_model": "N/A",
            "tokens_used": 0,
        })
        return {**state, "parsed_resume": {}, "raw_text": "", "errors": errors, "audit_entries": audit_entries}
