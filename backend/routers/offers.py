import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, HTTPException, Depends, BackgroundTasks
from fastapi.responses import StreamingResponse
from datetime import datetime
from bson import ObjectId
from database import db
from auth import require_role
from services.scheduler import send_offer_email

def serialize_doc(doc):
    if not doc:
        return None
    doc["id"] = str(doc["_id"])
    del doc["_id"]
    return doc

router = APIRouter(prefix="/offers", tags=["Offer Letter Agent"])

STANDARD_TEMPLATE = """HIREFLOW OFFICIAL OFFER OF EMPLOYMENT

Date: {date}

Dear {candidate_name},

We are thrilled to offer you the position of {job_title} at HireFlow Technologies. Based on your exceptional technical qualifications, interview performance, and project accomplishments, our leadership team believes you will make a profound contribution to our product mission.

OFFER TERMS & SUMMARY:
• Position Title: {job_title}
• Annual Compensation: {salary}
• Employment Type: Full-Time Regular
• Anticipated Start Date: {joining_date}
• Work Location: Remote / HQ Hybrid

BENEFITS & PERKS:
- Comprehensive Health, Vision, and Dental Insurance
- Flexible Paid Time Off (PTO) & Remote Workspace Allowance
- Continuous AI Learning & Conference Stipend

Please review this document and indicate your acceptance by signing via your HireFlow candidate portal by {joining_date}.

Sincerely,
HireFlow Talent Acquisition Team
HireFlow Technologies Inc.
"""

EXECUTIVE_TEMPLATE = """EXECUTIVE EMPLOYMENT AGREEMENT & OFFER LETTER

Date: {date}

Dear {candidate_name},

On behalf of HireFlow Technologies, I am delighted to extend an offer for the position of {job_title}. 

KEY COMPENSATION & TERMS:
1. Role: {job_title}
2. Compensation: Base Salary of {salary} per annum, paid semi-monthly.
3. Target Joining Date: {joining_date}
4. Equity & Bonus: Performance-based annual bonus + Stock options vesting over 4 years with a 1-year cliff.

CONFIDENTIALITY & GOVERNANCE:
This offer is contingent upon successful reference validation and standard employment background verification.

We are confident that your expertise will be instrumental in scaling HireFlow. We look forward to welcoming you aboard!

Warm regards,
Board of Directors & Executive Recruiting Team
HireFlow Technologies
"""

CONTRACTOR_TEMPLATE = """INDEPENDENT CONTRACTOR & CONSULTING OFFER

Date: {date}

Hello {candidate_name},

HireFlow is pleased to offer you an engagement as a Consultant for the position of {job_title}.

ENGAGEMENT HIGHLIGHTS:
• Engagement Title: {job_title}
• Agreed Compensation Rate: {salary}
• Project Effective Date: {joining_date}
• Contract Duration: 12 Months (Renewable)

DELIVERABLES & SCOPE:
You will collaborate directly with our engineering leadership to architect, build, and deploy agentic AI pipelines.

Please confirm your availability and formal acceptance via your candidate portal.

Regards,
Resource Management Team
HireFlow Technologies
"""

SAMPLE_TEMPLATES = [
    {
        "id": "standard",
        "title": "Standard Corporate Offer Letter",
        "template_text": STANDARD_TEMPLATE
    },
    {
        "id": "executive",
        "title": "Executive & Senior Engineering Agreement",
        "template_text": EXECUTIVE_TEMPLATE
    },
    {
        "id": "contractor",
        "title": "Contractor & Consulting Offer",
        "template_text": CONTRACTOR_TEMPLATE
    }
]

@router.get("/templates")
async def get_templates(current_user: dict = Depends(require_role(["hr"]))):
    """Retrieve sample offer letter templates."""
    return SAMPLE_TEMPLATES

@router.post("/generate/{application_id}")
async def generate_offer(
    application_id: str,
    payload: dict,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(require_role(["hr"]))
):
    """
    Generate a customized offer letter for a shortlisted application,
    updates application status to 'offered' in candidate portal, and dispatches email via Gmail SMTP.
    """
    if not ObjectId.is_valid(application_id):
        raise HTTPException(status_code=400, detail="Invalid application ID format")

    app_doc = await db.applications.find_one({"_id": ObjectId(application_id)})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    cand_email = app_doc["candidate_id"]
    user_doc = await db.users.find_one({"email": cand_email})
    cand_name = (
        (user_doc.get("name") if user_doc else None) or
        (user_doc.get("full_name") if user_doc else None) or
        app_doc.get("candidate_name") or
        cand_email.split("@")[0].replace(".", " ").capitalize()
    )

    job_title = "Software Engineer"
    if ObjectId.is_valid(app_doc["job_id"]):
        job_doc = await db.jobs.find_one({"_id": ObjectId(app_doc["job_id"])})
        if job_doc:
            job_title = job_doc.get("title", job_title)

    salary = payload.get("salary", "$120,000 / year")
    joining_date = payload.get("joining_date", "1st of Next Month")

    offer_body = DEFAULT_TEMPLATE.format(
        date=datetime.utcnow().strftime("%d %b %Y"),
        candidate_name=cand_name,
        job_title=job_title,
        salary=salary,
        joining_date=joining_date
    )

    now_iso = datetime.utcnow().isoformat()
    offer_doc = {
        "application_id": application_id,
        "candidate_email": cand_email,
        "candidate_name": cand_name,
        "job_id": app_doc["job_id"],
        "job_title": job_title,
        "salary": salary,
        "joining_date": joining_date,
        "offer_body": offer_body,
        "status": "sent",
        "created_by": current_user["email"],
        "created_at": now_iso,
        "updated_at": now_iso
    }

    # Save to offer_letters collection
    res = await db.offer_letters.insert_one(offer_doc)
    created = await db.offer_letters.find_one({"_id": res.inserted_id})

    # Update candidate application status to 'offered' so it displays in Candidate Portal
    await db.applications.update_one(
        {"_id": ObjectId(application_id)},
        {"$set": {"status": "offered", "updated_at": now_iso}}
    )

    # Background task: Send beautiful HTML email to candidate's respective email address
    background_tasks.add_task(
        send_offer_email,
        candidate_email=cand_email,
        candidate_name=cand_name,
        job_title=job_title,
        salary=salary,
        joining_date=joining_date,
        offer_body=offer_body
    )

    return serialize_doc(created)

@router.post("/send/{offer_id}")
async def send_existing_offer(
    offer_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(require_role(["hr"]))
):
    """Dispatch an existing offer letter to the candidate's respective email address."""
    if not ObjectId.is_valid(offer_id):
        raise HTTPException(status_code=400, detail="Invalid offer ID format")

    offer = await db.offer_letters.find_one({"_id": ObjectId(offer_id)})
    if not offer:
        raise HTTPException(status_code=404, detail="Offer letter not found")

    # Update application status to 'offered'
    if ObjectId.is_valid(offer["application_id"]):
        await db.applications.update_one(
            {"_id": ObjectId(offer["application_id"])},
            {"$set": {"status": "offered", "updated_at": datetime.utcnow().isoformat()}}
        )

    # Update offer letter status
    await db.offer_letters.update_one(
        {"_id": ObjectId(offer_id)},
        {"$set": {"status": "sent", "updated_at": datetime.utcnow().isoformat()}}
    )

    background_tasks.add_task(
        send_offer_email,
        candidate_email=offer["candidate_email"],
        candidate_name=offer["candidate_name"],
        job_title=offer["job_title"],
        salary=offer["salary"],
        joining_date=offer["joining_date"],
        offer_body=offer["offer_body"]
    )

    return {"message": f"Offer letter successfully dispatched to {offer['candidate_email']}"}

@router.get("/list")
async def list_offers(current_user: dict = Depends(require_role(["hr"]))):
    """List all generated offer letters."""
    cursor = db.offer_letters.find().sort("created_at", -1)
    offers = []
    async for doc in cursor:
        offers.append(serialize_doc(doc))
    return offers

@router.patch("/{offer_id}")
async def update_offer(
    offer_id: str,
    payload: dict,
    current_user: dict = Depends(require_role(["hr"]))
):
    """Edit generated offer letter details or body."""
    if not ObjectId.is_valid(offer_id):
        raise HTTPException(status_code=400, detail="Invalid offer ID format")

    update_fields = {}
    if "offer_body" in payload:
        update_fields["offer_body"] = payload["offer_body"]
    if "salary" in payload:
        update_fields["salary"] = payload["salary"]
    if "joining_date" in payload:
        update_fields["joining_date"] = payload["joining_date"]

    update_fields["updated_at"] = datetime.utcnow().isoformat()

    await db.offer_letters.update_one({"_id": ObjectId(offer_id)}, {"$set": update_fields})
    updated = await db.offer_letters.find_one({"_id": ObjectId(offer_id)})
    return serialize_doc(updated)

@router.get("/{offer_id}/download")
async def download_offer(
    offer_id: str,
    current_user: dict = Depends(require_role(["hr"]))
):
    """Download offer letter text content."""
    if not ObjectId.is_valid(offer_id):
        raise HTTPException(status_code=400, detail="Invalid offer ID format")

    offer = await db.offer_letters.find_one({"_id": ObjectId(offer_id)})
    if not offer:
        raise HTTPException(status_code=404, detail="Offer letter not found")

    return {
        "id": offer_id,
        "candidate_name": offer.get("candidate_name"),
        "filename": f"Offer_Letter_{offer.get('candidate_name','Candidate').replace(' ','_')}.txt",
        "content": offer.get("offer_body")
    }

@router.get("/{offer_id}/download-docx")
async def download_offer_docx(
    offer_id: str,
    current_user: dict = Depends(require_role(["hr"]))
):
    """Download styled Word Document (.docx) for the offer letter."""
    if not ObjectId.is_valid(offer_id):
        raise HTTPException(status_code=400, detail="Invalid offer ID format")

    offer = await db.offer_letters.find_one({"_id": ObjectId(offer_id)})
    if not offer:
        raise HTTPException(status_code=404, detail="Offer letter not found")

    doc = Document()
    
    # Title Header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = header_p.add_run("HIREFLOW TECHNOLOGIES\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229)

    run_sub = header_p.add_run("OFFICIAL EMPLOYMENT OFFER\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacing

    # Content paragraphs
    offer_body = offer.get("offer_body", "")
    for line in offer_body.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # Save docx into memory stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Offer_Letter_{offer.get('candidate_name','Candidate').replace(' ', '_')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
